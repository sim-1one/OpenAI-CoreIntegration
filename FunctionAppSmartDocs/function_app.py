import os
import logging
import docx
import csv
import requests
import collections.abc
import azure.functions as func
from azure.identity import ManagedIdentityCredential
from azure.mgmt.resource import ResourceManagementClient
from azure.mgmt.subscription import SubscriptionClient

# Configurazione logging avanzato
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Endpoint OpenAI
QUESTION_ENDPOINT = os.getenv("QUESTION_ENDPOINT")
logging.info(f"Question endpoint set to: {QUESTION_ENDPOINT}")

# Autenticazione tramite identità gestita della Function App
credential = ManagedIdentityCredential()
subscription_client = SubscriptionClient(credential)
tokenai = credential.get_token("https://cognitiveservices.azure.com/.default")


# Funzione per appiattire un dizionario annidato in una struttura piana
def flatten_dict(d, parent_key='', sep='_'):
    #logging.info(f"start function flatten_dict")
    items = []
    for k, v in d.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k

        if isinstance(v, collections.abc.MutableMapping):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        elif isinstance(v, list):
            for i, item in enumerate(v):
                if isinstance(item, collections.abc.MutableMapping):
                    items.extend(flatten_dict(item, f"{new_key}{sep}{i}", sep=sep).items())
                else:
                    items.append((f"{new_key}{sep}{i}", item))
        else:
            items.append((new_key, v))
    return dict(items)

# Funzione per ottenere tutte le risorse con un determinato tag in una sottoscrizione
def get_resources_by_tag_in_subscription(subscription_id, tag_key, tag_value):
    logging.info(f"start get_resources_by_tag_in_subscription")
    resource_client = ResourceManagementClient(credential, subscription_id)
    tag_filter = f"tagName eq '{tag_key}' and tagValue eq '{tag_value}'"
    logging.info(f"Cercando risorse con il tag: {tag_key}={tag_value} nella sottoscrizione {subscription_id}")

    resources = resource_client.resources.list(filter=tag_filter)
    resources_list = []

    for resource in resources:
        resource_details = {
            'name': resource.name,
            'id': resource.id,
            'location': resource.location,
            'type': resource.type,
            'tags': resource.tags
        }
        resources_list.append(resource_details)
        logging.info(f"Trovata risorsa: {resource.name} - {resource.type}")

    logging.info(str(resources_list))
    return resources_list

# Funzione per ottenere i Resource Groups con un determinato tag in una sottoscrizione
def get_resource_groups_by_tag_in_subscription(subscription_id, tag_key, tag_value):
    logging.info(f"start get_resource_groups_by_tag_in_subscription")
    resource_client = ResourceManagementClient(credential, subscription_id)
    logging.info(f"Cercando resource groups con il tag: {tag_key}={tag_value} nella sottoscrizione {subscription_id}")
    resource_groups = resource_client.resource_groups.list()
    matching_resource_groups = []

    for rg in resource_groups:
        if rg.tags and tag_key in rg.tags and rg.tags[tag_key] == tag_value:
            logging.info(f"Trovato resource group: {rg.name}")
            matching_resource_groups.append(rg)
    logging.info(str(matching_resource_groups))
    return matching_resource_groups

# Funzione per ottenere tutte le risorse con un determinato tag da tutte le sottoscrizioni
def get_all_resources(tag_key, tag_value):
    logging.info(f"start get_all_resources")
    all_resources = []
    added_resource_ids = set()

    # Elenca tutte le sottoscrizioni a cui hai accesso
    for subscription in subscription_client.subscriptions.list():
        subscription_id = subscription.subscription_id
        logging.info(f"Processando la sottoscrizione: {subscription_id}")

        # Cerca le risorse con il tag specificato in questa sottoscrizione
        resources = get_resources_by_tag_in_subscription(subscription_id, tag_key, tag_value)
        for resource in resources:
            logging.info(f"Analizzando risorsa {resource}")
            if resource['id'] not in added_resource_ids:
                all_resources.append(resource)
                added_resource_ids.add(resource['id'])

        # Cerca i resource group con il tag specificato in questa sottoscrizione
        resource_groups = get_resource_groups_by_tag_in_subscription(subscription_id, tag_key, tag_value)
        for rg in resource_groups:
            logging.info(f"Analizzando {rg}")
            # Cerca risorse all'interno di ciascun resource group
            rg_resources = get_resources_in_resource_group_in_subscription(subscription_id, rg.name)
            for resource in rg_resources:
                if resource['id'] not in added_resource_ids:
                    all_resources.append(resource)
                    added_resource_ids.add(resource['id'])
    logging.info(str(all_resources))
    return all_resources

# Funzione per ottenere le risorse all'interno di un Resource Group in una sottoscrizione
def get_resources_in_resource_group_in_subscription(subscription_id, resource_group_name):
    logging.info(f"start get_resources_in_resource_group_in_subscription")
    resource_client = ResourceManagementClient(credential, subscription_id)
    logging.info(f"Recuperando risorse dal resource group: {resource_group_name} nella sottoscrizione {subscription_id}")
    resources = resource_client.resources.list_by_resource_group(resource_group_name)
    resources_list = []

    for resource in resources:
        resource_details = {
            'name': resource.name,
            'id': resource.id,
            'location': resource.location,
            'type': resource.type,
            'tags': resource.tags
        }
        resources_list.append(resource_details)
        logging.info(f"Trovata risorsa nel resource group {resource_group_name}: {resource.name} - {resource.type}")
    logging.info(str(resources_list))
    return resources_list

# Funzione per ottenere l'API più recente per una risorsa specifica
def get_latest_api_version(resource_client, resource_type):
    logging.info(f"start get_latest_api_version")
    provider_namespace, resource_type_name = resource_type.split('/', 1)
    provider = resource_client.providers.get(provider_namespace)
    resource_type_info = next(
        (t for t in provider.resource_types if t.resource_type == resource_type_name), None
    )
    if resource_type_info:
        return sorted(resource_type_info.api_versions, reverse=True)[0]
    return None

# Funzione per ottenere i metadati completi di una risorsa specifica
def get_resource_metadata(resource_client, resource):
    logging.info(f"start get_resource_metadata")
    resource_type = resource['type']
    api_version = get_latest_api_version(resource_client, resource_type)
    if api_version:
        logging.info(f"Usando l'API version: {api_version} per la risorsa {resource['name']}")
        resource_metadata = resource_client.resources.get_by_id(resource['id'], api_version=api_version)
        return resource_metadata
    else:
        logging.info(f"Impossibile trovare l'API per la risorsa {resource['name']} con tipo {resource['type']}")
        return None

# Funzione per generare la overview del workload leggendo il file CSV
def generate_workload_overview():
    logging.info(f"start generate_workload_overview")
    resources_info = []

    with open("/tmp/resources_with_expanded_metadata.csv", 'r', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            resources_info.append(row)

    resources_str = "\n".join([
        f"Name: {row['Name']}, Type: {row['Type']}, Location: {row['Location']}, Resource Group: {row['Resource ID'].split('/')[4]}"
        for row in resources_info])

    payload = {
        "messages": [
            {"role": "system",
             "content": "You are an expert Azure Architect and Documentation Writer. Your job is to create a clear and detailed overview of an Azure workload."},
            {"role": "user",
             "content": f"Here is the list of resources in the workload:\n{resources_str}.\nGenerate a detailed and human-readable overview. Explain at the end how the workload is splitted, so if is multi-regional or not and so on."}
        ],
        "temperature": 0.7,
        "max_tokens": 16000
    }

    response = requests.post(
        QUESTION_ENDPOINT,
        headers={'content-type': 'application/json', 'Authorization': 'Bearer ' + tokenai.token},
        json=payload
    )

    response.raise_for_status()
    response_from_copilot = response.json()['choices'][0]['message']['content'].strip()
    logging.info(str(response_from_copilot))
    return response_from_copilot

# Funzione per generare la documentazione con OpenAI
def generate_infra_config(metadata_list):
    logging.info("Inizio della generazione della configurazione dell'infrastruttura.")
    document_content = ""

    for index, metadata in enumerate(metadata_list):
        logging.debug(f"Elaborazione del metadato {index + 1}/{len(metadata_list)}: {metadata}")
        metadata_str = str(metadata)

        payload = {
            "messages": [
                {"role": "system", "content": "You are an expert Azure Architect and Documentation Writer."},
                {"role": "user", "content": f"Here is the metadata for an Azure resource: \n{metadata_str}.\nPlease generate a detailed and human-readable documentation."}
            ],
            "temperature": 0.7,
            "max_tokens": 16000
        }

        try:
            # Invio richiesta a OpenAI
            logging.debug("Invio della richiesta a OpenAI con payload.")
            response = requests.post(
                QUESTION_ENDPOINT,
                headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {tokenai.token}'},
                json=payload
            )

            # Gestione della risposta
            response.raise_for_status()
            response_data = response.json()
            response_from_copilot = response_data['choices'][0]['message']['content'].strip()
            logging.info(f"Risposta da OpenAI per il metadato {index + 1}: {response_from_copilot}")

        except requests.exceptions.RequestException as e:
            logging.error(f"Errore nella richiesta a OpenAI per il metadato {index + 1}: {e}")
            continue  # Salta al prossimo elemento se c'è un errore

        # Revisione del contenuto tramite ArchitecturalReviewer e DocCreator
        response_from_DocCreator = response_from_copilot
        previousdoc = response_from_DocCreator
        for i in range(3):
            try:
                ArchitecturalReviewer_response = ArchitecturalReviewer(response_from_DocCreator)
                logging.debug(f"Risposta ArchitecturalReviewer (ciclo {i + 1}): {ArchitecturalReviewer_response}")
                response_from_DocCreator, previousdoc = DocCreator(ArchitecturalReviewer_response, previousdoc)
                logging.debug(f"Documento revisionato (ciclo {i + 1}): {response_from_DocCreator}")
            except Exception as e:
                logging.error(f"Errore durante la revisione della documentazione (ciclo {i + 1}): {e}")
                break  # Interrompe la revisione in caso di errore

        # Aggiunge la documentazione revisionata al contenuto finale
        document_content += response_from_DocCreator + "\n\n"

    # Scrittura del contenuto finale in un file
    try:
        logging.info("Scrittura del documento finale nel file architecture.txt.")
        with open("/tmp/architecture.txt", "a", encoding="utf-8") as architecturefile:
            architecturefile.write(document_content)
        logging.info("Documento architecture.txt generato con successo.")
    except IOError as e:
        logging.error(f"Errore durante la scrittura del file architecture.txt: {e}")

    return document_content

    logging.info(f"start generate_infra_config")
    document_content = ""

    for metadata in metadata_list:
        metadata_str = str(metadata)

        payload = {
            "messages": [
                {"role": "system",
                 "content": "You are an expert Azure Architect and Documentation Writer."},
                {"role": "user",
                 "content": f"Here is the metadata for an Azure resource: \n{metadata_str}.\nPlease generate a detailed and human-readable documentation."}
            ],
            "temperature": 0.7,
            "max_tokens": 16000
        }

        response = requests.post(
            QUESTION_ENDPOINT,
            headers={'content-type': 'application/json', 'Authorization': 'Bearer ' + tokenai.token},
            json=payload
        )

        response.raise_for_status()
        response_from_copilot = response.json()['choices'][0]['message']['content'].strip()
        logging.info(str(response_from_copilot))

        # Passa il contenuto attraverso 3 cicli di revisione con ArchitecturalReviewer e DocCreator
        response_from_DocCreator = response_from_copilot
        previousdoc = response_from_DocCreator
        for i in range(3):
            ArchitecturalReviewer_response = ArchitecturalReviewer(response_from_DocCreator)
            response_from_DocCreator, previousdoc = DocCreator(ArchitecturalReviewer_response, previousdoc)

        # Aggiunge la documentazione revisionata al contenuto finale
        document_content += response_from_DocCreator + "\n\n"

    with open("architecture.txt", "a", encoding="utf-8") as architecturefile:
        architecturefile.write(document_content)

    return document_content

# Funzione per convertire il file txt in docx e aggiungere l'overview
def txt_to_docx():
    logging.info(f"start txt_to_docx")
    logging.info("Generazione dei file in corso...")

    doc = docx.Document()

    # Genera l'overview del workload
    overview = generate_workload_overview()

    # Aggiungi "Workload Overview" come Titolo 1
    doc.add_heading("Workload Overview", level=1)
    doc.add_paragraph(overview)

    # Aggiungi un'interruzione di pagina per iniziare i dettagli su una nuova pagina
    doc.add_page_break()

    # Aggiungi "Workload Details" come Titolo 1
    doc.add_heading("Workload Details", level=1)

    # Aggiungi il contenuto del file architecture.txt
    with open("architecture.txt", 'r', encoding='utf-8', errors='ignore') as openfile:
        line = openfile.read()
        doc.add_paragraph(line)

    # Salva il documento Word
    doc.save("Output.docx")
    logging.info("Il file Output.docx è stato creato con successo.")

def cleanup_files():
    logging.info(f"start cleanup_files")
    logging.info("Cleaning up temporary files")
    if os.path.exists("architecture.txt"):
        os.remove("architecture.txt")
        logging.debug("architecture.txt deleted")
    if os.path.exists("resources_with_expanded_metadata.csv"):
        os.remove("resources_with_expanded_metadata.csv")
        logging.debug("resources_with_expanded_metadata.csv")
    if os.path.exists("Output.docx"):
        os.remove("Output.docx")
        logging.debug("Output.docx deleted")
    if os.path.exists("/tmp/architecture.txt"):
        os.remove("/tmp/architecture.txt")
        logging.debug("architecture.txt deleted")
    if os.path.exists("/tmp/resources_with_expanded_metadata.csv"):
        os.remove("/tmp/resources_with_expanded_metadata.csv")
        logging.debug("resources_with_expanded_metadata.csv")
    if os.path.exists("/tmp/Output.docx"):
        os.remove("/tmp/Output.docx")
        logging.debug("Output.docx deleted")
    logging.info("Temporary files cleanup completed")


def save_resources_with_expanded_metadata_to_csv(resources, metadata_list):
    logging.info("Start save_resources_with_expanded_metadata_to_csv")
    
    all_keys = set()

    # Raccogli tutte le chiavi disponibili nei metadati, incluse quelle dei dizionari annidati
    try:
        for metadata in metadata_list:
            flat_metadata = flatten_dict(metadata.__dict__)  # Appiattiamo il dizionario dei metadati
            logging.debug(f"Flat metadata for item: {flat_metadata}")
            all_keys.update(flat_metadata.keys())
    except AttributeError as e:
        logging.error(f"Errore durante l'appiattimento dei metadati: {e}")
        return
    except Exception as e:
        logging.error(f"Errore sconosciuto durante la raccolta delle chiavi: {e}")
        return

    # Trasforma il set delle chiavi in una lista ordinata per mantenere ordine nelle colonne
    all_keys = list(all_keys)
    logging.info(f"All unique metadata keys: {all_keys}")

    # Scrivi le risorse e i loro metadati in un file CSV
    try:
        with open("/tmp/resources_with_expanded_metadata.csv", mode="w", newline='', encoding="utf-8") as csv_file:
            # Creiamo l'header del CSV con tutte le chiavi uniche
            fieldnames = ['Name', 'Resource ID', 'Location', 'Type', 'Tags'] + all_keys
            logging.debug(f"Fieldnames for CSV: {fieldnames}")
            writer = csv.DictWriter(csv_file, fieldnames=fieldnames)

            # Scriviamo l'intestazione
            writer.writeheader()
            logging.info("CSV header written successfully.")

            # Scrivi i dettagli di ogni risorsa e i suoi metadati
            for resource, metadata in zip(resources, metadata_list):
                logging.debug(f"Processing resource: {resource['name']}")

                # Crea la riga iniziale per la risorsa
                resource_row = {
                    'Name': resource['name'],
                    'Resource ID': resource['id'],
                    'Location': resource['location'],
                    'Type': resource['type'],
                    'Tags': resource['tags']
                }

                # Appiattiamo i metadati annidati prima di scriverli nel CSV
                try:
                    flat_metadata = flatten_dict(metadata.__dict__)
                    logging.debug(f"Flat metadata for resource {resource['name']}: {flat_metadata}")
                except AttributeError as e:
                    logging.error(f"Errore durante l'appiattimento dei metadati per la risorsa {resource['name']}: {e}")
                    continue
                except Exception as e:
                    logging.error(f"Errore sconosciuto durante l'appiattimento dei metadati per {resource['name']}: {e}")
                    continue

                # Aggiungi i metadati alle colonne, gestendo eventuali valori mancanti
                for key in all_keys:
                    resource_row[key] = flat_metadata.get(key, 'N/A')  # Usa 'N/A' per valori mancanti

                # Scrivi la riga nel CSV
                writer.writerow(resource_row)
                logging.debug(f"Row written for resource {resource['name']}")

    except IOError as e:
        logging.error(f"Errore durante la creazione del file CSV: {e}")
    except Exception as e:
        logging.error(f"Errore sconosciuto durante la scrittura nel file CSV: {e}")

    logging.info("Il file resources_with_expanded_metadata.csv è stato creato con successo.")

    logging.info(f"start save_resources_with_expanded_metadata_to_csv")
    all_keys = set()

    # Raccogli tutte le chiavi disponibili nei metadati, incluse quelle dei dizionari annidati
    for metadata in metadata_list:
        flat_metadata = flatten_dict(metadata.__dict__)  # Appiattiamo il dizionario dei metadati
        all_keys.update(flat_metadata.keys())

    # Trasforma il set delle chiavi in una lista ordinata per mantenere ordine nelle colonne
    all_keys = list(all_keys)
    logging.info(f"all_keys: {all_keys}")

    # Scrivi le risorse e i loro metadati in un file CSV
    with open("/tmp/resources_with_expanded_metadata.csv", mode="w", newline='', encoding="utf-8") as csv_file:
        # Creiamo l'header del CSV con tutte le chiavi uniche
        fieldnames = ['Name', 'Resource ID', 'Location', 'Type', 'Tags'] + all_keys
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames)

        # Scriviamo l'intestazione
        writer.writeheader()

        # Scrivi i dettagli di ogni risorsa e i suoi metadati
        for resource, metadata in zip(resources, metadata_list):
            resource_row = {
                'Name': resource['name'],
                'Resource ID': resource['id'],
                'Location': resource['location'],
                'Type': resource['type'],
                'Tags': resource['tags']
            }

            # Appiattiamo i metadati annidati prima di scriverli nel CSV
            flat_metadata = flatten_dict(metadata.__dict__)

            # Aggiungi i metadati alle colonne, gestendo eventuali valori mancanti
            for key in all_keys:
                resource_row[key] = flat_metadata.get(key, 'N/A')  # Usa 'N/A' per valori mancanti

            # Scrivi la riga nel CSV
            writer.writerow(resource_row)

    logging.info("Il file resources_with_expanded_metadata.csv è stato creato con successo.")

def ArchitecturalReviewer(response_from_DocCreator):
    logging.info(f"start ArchitecturalReviewer")
    payload = {
        "messages": [
            {"role": "system",
             "content": "You are the Azure architectural reviewer of the enterprise. Our team is creating documentation on our Azure architecture for the existing dedicated workload. The team will pass a specific piece of documentation to you each time. Make suggestions on how to make it more user-friendly the part without suggesting to add graph, diagrams, table of contents, feedback and so on. Suggest how to write the doc in a user-friendly readable way."},
            {"role": "user",
             "content": f"{response_from_DocCreator}"}
        ],
        "temperature": 0.7,
        "max_tokens": 16000
    }

    response = requests.post(
        QUESTION_ENDPOINT,
        headers={'content-type': 'application/json', 'Authorization': 'Bearer ' + tokenai.token},
        json=payload
    )

    response.raise_for_status()
    ArchitecturalReviewer_response = response.json()['choices'][0]['message']['content'].strip()
    logging.info(f"Architectural Reviewer Comments: {ArchitecturalReviewer_response}")

    return ArchitecturalReviewer_response

def DocCreator(ArchitecturalReviewer_response,previousdoc):
    logging.info(f"start DocCreator")
    payload = {
        "messages": [
            {"role": "system",
             "content": "You have created a document about your Azure infrastructure related a workload. Your supervisor is reviewing the documentation. Generate a new documentation output based on his suggestions as an output."},
            {"role": "user",
             "content": f"source:{previousdoc}. Suggestion: {ArchitecturalReviewer_response}"}
        ],
        "temperature": 0.7,
        "max_tokens": 16000
    }

    response = requests.post(
        QUESTION_ENDPOINT,
        headers={'content-type': 'application/json', 'Authorization': 'Bearer ' + tokenai.token},
        json=payload
    )

    response.raise_for_status()
    response_from_DocCreator = response.json()['choices'][0]['message']['content'].strip()
    logging.info(f"Doc Creator Repsonse: {response_from_DocCreator}")

    return response_from_DocCreator, response_from_DocCreator

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

@app.route(route="smartdocs")
def smartdocs(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("Received request for smartdocs function")

    tag_key = req.params.get("tag_key")
    tag_value = req.params.get("tag_value")
    if not tag_key or not tag_value:
        try:
            req_body = req.get_json()
            tag_key = tag_key or req_body.get("tag_key")
            tag_value = tag_value or req_body.get("tag_value")
        except ValueError:
            logging.error("Invalid request body")
            return func.HttpResponse("Missing 'tag_key' or 'tag_value' parameters.", status_code=400)

    if not tag_key or not tag_value:
        logging.error("Missing required parameters: 'tag_key' and 'tag_value'")
        return func.HttpResponse("Missing 'tag_key' or 'tag_value' parameters.", status_code=400)

    all_resources = get_all_resources(tag_key, tag_value)
    metadata_list = []
    processed_resource_ids = set()

    for subscription in subscription_client.subscriptions.list():
        subscription_id = subscription.subscription_id
        resource_client = ResourceManagementClient(credential, subscription_id)
        for resource in all_resources:
            if resource['id'] not in processed_resource_ids:
                metadata = get_resource_metadata(resource_client, resource)
                if metadata:
                    metadata_list.append(metadata)
                    processed_resource_ids.add(resource['id'])

    save_resources_with_expanded_metadata_to_csv(all_resources, metadata_list)
    document_content = generate_infra_config(metadata_list)

    with open("architecture.txt", "w", encoding="utf-8") as output_file:
        output_file.write(document_content)
        logging.info("Architecture text file saved")

    txt_to_docx()
    cleanup_files()

    logging.info("Function smartdocs completed successfully")
    return func.HttpResponse("Resource data processed successfully.", status_code=200)
